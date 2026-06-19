"""
report_docx.py
EFARS 보고서(별지 제2호서식 약식)를 writer JSON 으로부터 .docx 로 렌더링한다.

방식: backend/templates/별지_제2호서식_템플릿_약식.docx 를 열어
      {placeholder} 토큰을 writer JSON 값으로 치환한다.
      (코드 오케스트레이터는 Code Interpreter 미가용이라 에이전트가 파일을 못 만들므로
       백엔드가 동일 템플릿으로 직접 렌더링)

템플릿 파일은 반드시 레포에 포함되어야 한다(에이전트 업로드본은 깃 백엔드가 접근 불가).
템플릿이 없으면 from-scratch 레이아웃으로 안전하게 폴백한다.

사용:
    from report_docx import build_report_docx
    docx_bytes = build_report_docx(analyze_data)   # /analyze 응답의 data 딕셔너리
"""

import os
from io import BytesIO
from datetime import datetime
from typing import Any, Dict, List, Optional

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

KR_FONT = "맑은 고딕"
ACCENT = RGBColor(0x1F, 0x4E, 0x79)
MUTED = RGBColor(0x55, 0x5F, 0x70)
PLACEHOLDER = "[데이터 미확보]"

TEMPLATE_PATH = os.path.join(
    os.path.dirname(os.path.abspath(__file__)),
    "templates",
    "별지_제2호서식_템플릿_약식.docx",
)


# ──────────────────────────────────────────────────────────────────────────
# 공통 폰트 헬퍼
# ──────────────────────────────────────────────────────────────────────────
def _apply_kr_font(run, name: str = KR_FONT) -> None:
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), name)


def _g(d: Any, *keys, default="") -> Any:
    cur = d
    for k in keys:
        if not isinstance(cur, dict):
            return default
        cur = cur.get(k, default)
    return cur if cur not in (None, "") else default


def _nz(value: Any) -> str:
    """빈 값이면 [데이터 미확보]."""
    s = "" if value is None else str(value).strip()
    return s if s else PLACEHOLDER


# ──────────────────────────────────────────────────────────────────────────
# 값 매핑 빌더
# ──────────────────────────────────────────────────────────────────────────
def _build_mapping(data: Dict[str, Any]) -> Dict[str, str]:
    classification = data.get("classification", {}) or {}
    regulations = data.get("regulations", {}) or {}
    timeline = data.get("timeline", {}) or {}
    wr = data.get("report_writer", {}) or {}
    _es_raw = wr.get("executive_summary", {})
    # executive_summary 가 문자열로 올 경우 dict 로 강제 변환 방지
    es = _es_raw if isinstance(_es_raw, dict) else {}
    rm = wr.get("responsibility_matrix", {}) or {}
    agent_inputs = data.get("agent_inputs", {}) or {}

    case_id = data.get("case_id", "") or "EFARS"
    t0 = _g(agent_inputs, "code-orchestrator", "t0_kst") or _g(es, "timing")

    sev = _g(classification, "severity_class")
    level = _g(classification, "severity_response_level")
    sev_combined = f"{sev} · {level}".strip(" ·") if (sev or level) else PLACEHOLDER

    # 임원요약 (멀티라인)
    if es:
        summary_lines = [
            f'사건: {_nz(es.get("incident"))}',
            f'시점: {_nz(es.get("timing"))}',
            f'유형·등급: {_nz(es.get("incident_type"))} / {_nz(sev_combined)}',
            f'등급 확정 권한: {_nz(es.get("severity_decision_authority"))}',
            f'현재 상태: {_nz(es.get("current_status"))}',
        ]
        kd = es.get("key_decisions") or []
        if kd:
            summary_lines.append(f'결정 요청: {", ".join(str(x) for x in kd)}')
        exec_summary = "\n".join(summary_lines)
    else:
        # executive_summary 가 문자열인 경우 그대로 사용
        _es_str = wr.get("executive_summary", "")
        if isinstance(_es_str, str) and _es_str.strip():
            exec_summary = _es_str.strip()
        else:
            body = wr.get("report_markdown") or ""
            exec_summary = "\n".join(l.strip() for l in str(body).splitlines() if l.strip()) or PLACEHOLDER

    # 주요근거조항
    clauses = regulations.get("applicable_clauses") if isinstance(regulations, dict) else None
    if isinstance(clauses, list) and clauses:
        picked = []
        internal = [c for c in clauses if isinstance(c, dict) and "사내" in str(c.get("source", ""))]
        ordered = (internal or clauses)[:2]
        for c in ordered:
            if isinstance(c, dict):
                picked.append(f'{c.get("source", "")} {c.get("article", "") or c.get("title", "")}'.strip())
            else:
                picked.append(str(c))
        main_basis = " · ".join(p for p in picked if p) or PLACEHOLDER
    else:
        main_basis = PLACEHOLDER

    # 결정요청 (멀티라인, 최대 3)
    decisions = wr.get("decisions_requested", []) or []
    if decisions:
        d_lines = []
        for d in decisions[:3]:
            if isinstance(d, dict):
                d_lines.append(f'- {d.get("decision", "")} · 권고: {d.get("recommendation", "-")}')
            else:
                d_lines.append(f"- {d}")
        decisions_txt = "\n".join(d_lines)
    else:
        decisions_txt = PLACEHOLDER

    # 후속조치 (멀티라인, 최대 3)
    follow_ups = wr.get("follow_up_actions", []) or []
    if follow_ups:
        f_lines = []
        for a in follow_ups[:3]:
            if isinstance(a, dict):
                f_lines.append(
                    f'- {a.get("action", "")} '
                    f'(시한: {a.get("due_kst", "-")}, 담당: {a.get("responsible_role", "-")})'
                )
            else:
                f_lines.append(f"- {a}")
        follow_txt = "\n".join(f_lines)
    else:
        follow_txt = PLACEHOLDER

    missing = wr.get("missing_fields", []) or []
    missing_txt = ", ".join(str(m) for m in missing) if missing else "없음"

    mapping = {
        "보고일자": datetime.now().strftime("%Y-%m-%d %H:%M"),
        "보고종류": "1차 보고(약식)",
        "case_id": case_id,
        "T0_KST": _nz(t0),
        "사고유형": _nz(classification.get("incident_type")),
        "사고등급": _nz(sev_combined),
        "등급확정권한": _nz(classification.get("severity_decision_authority")),
        "보고기한_최초": _nz(timeline.get("initial_due")),
        "보고채널": "EFARS 전자제출",
        "임원요약": exec_summary,
        "주요근거조항": main_basis,
        "1차보고_비고": "-",
        "경과보고_비고": "-",
        "원인분석_비고": "-",
        "사고선포상태": _nz(wr.get("declaration_status")) if wr.get("declaration_status") else "미선포",
        "완료조치": _nz(wr.get("completed_actions")) if wr.get("completed_actions") else PLACEHOLDER,
        "대응조직": _nz(wr.get("response_org")) if wr.get("response_org") else "미선포",
        "총괄상태": "검토 대기",
        "실무상태": "작성 완료",
        "선포상태": "검토 대기",
        "결정요청": decisions_txt,
        "후속조치": follow_txt,
        "다음보고일정": _nz(timeline.get("initial_due")),
        "self_check_score": str(wr.get("self_check_score", 0)),
        "missing_fields_요약": missing_txt,
    }
    return mapping


# ──────────────────────────────────────────────────────────────────────────
# 치환 헬퍼 (run 분할/멀티라인 대응)
# ──────────────────────────────────────────────────────────────────────────
def _set_paragraph(paragraph, new_text: str) -> None:
    """단락 전체 텍스트를 new_text 로 교체. '\\n' 은 실제 줄바꿈으로. 첫 run 서식 보존."""
    runs = paragraph.runs
    font_name = KR_FONT
    font_size = None
    font_bold = None
    if runs:
        r0 = runs[0]
        font_name = r0.font.name or KR_FONT
        font_size = r0.font.size
        font_bold = r0.font.bold
    for r in list(runs):
        r._element.getparent().remove(r._element)

    lines = new_text.split("\n")
    for i, line in enumerate(lines):
        run = paragraph.add_run(line)
        if font_size is not None:
            run.font.size = font_size
        if font_bold is not None:
            run.font.bold = font_bold
        _apply_kr_font(run, font_name)
        if i < len(lines) - 1:
            run.add_break()


def _replace_placeholders_in_paragraph(paragraph, mapping: Dict[str, str]) -> None:
    full = "".join(r.text for r in paragraph.runs)
    if "{" not in full:
        return
    new = full
    for key, val in mapping.items():
        token = "{" + key + "}"
        if token in new:
            new = new.replace(token, val if val not in (None, "") else PLACEHOLDER)
    if new != full:
        _set_paragraph(paragraph, new)


def _set_cell(cell, text: str) -> None:
    # 첫 단락에 기록, 나머지 단락 제거
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    _set_paragraph(cell.paragraphs[0], text)


def _fill_timeline_table(doc: Document, data: Dict[str, Any]) -> None:
    """헤더가 '시각'으로 시작하는 3열 테이블의 데이터 행을 채운다."""
    _es_raw = (data.get("report_writer", {}) or {}).get("executive_summary", {})
    es = _es_raw if isinstance(_es_raw, dict) else {}
    timeline = data.get("timeline", {}) or {}
    agent_inputs = data.get("agent_inputs", {}) or {}
    t0 = _g(agent_inputs, "code-orchestrator", "t0_kst") or _g(es, "timing")

    for t in doc.tables:
        if not t.rows or len(t.columns) < 3:
            continue
        header = t.rows[0].cells[0].text.strip()
        if "시각" in header and len(t.rows) >= 2:
            drow = t.rows[1]
            # 병합 셀 감지: 동일 _tc 를 가리키는 셀은 1개로 취급
            unique = []
            seen = set()
            for c in drow.cells:
                key = id(c._tc)
                if key not in seen:
                    seen.add(key)
                    unique.append(c)
            time_v = _nz(t0 or timeline.get("initial_due"))
            event_v = _nz(es.get("incident"))
            if len(unique) >= 3:
                _set_cell(unique[0], time_v)
                _set_cell(unique[1], event_v)
                _set_cell(unique[2], PLACEHOLDER)
            else:
                # 병합된 단일 셀 → 한 줄로 합쳐 표기
                _set_cell(unique[0], f"{time_v}  ·  {event_v}")
            return


# ──────────────────────────────────────────────────────────────────────────
# 메인 엔트리
# ──────────────────────────────────────────────────────────────────────────
def build_report_docx(data: Dict[str, Any]) -> bytes:
    if os.path.exists(TEMPLATE_PATH):
        return _build_from_template(data)
    return _build_from_scratch(data)


def _build_from_template(data: Dict[str, Any]) -> bytes:
    doc = Document(TEMPLATE_PATH)
    mapping = _build_mapping(data)

    # 타임라인 테이블 먼저 개별 셀 처리(3칸이 모두 {타임라인}이라 일괄치환 부적합)
    _fill_timeline_table(doc, data)

    # 본문 단락 치환
    for p in doc.paragraphs:
        _replace_placeholders_in_paragraph(p, mapping)

    # 모든 테이블 셀 치환
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_placeholders_in_paragraph(p, mapping)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ──────────────────────────────────────────────────────────────────────────
# 폴백: 템플릿이 없을 때 from-scratch
# ──────────────────────────────────────────────────────────────────────────
def _build_from_scratch(data: Dict[str, Any]) -> bytes:
    classification = data.get("classification", {}) or {}
    timeline = data.get("timeline", {}) or {}
    wr = data.get("report_writer", {}) or {}
    case_id = data.get("case_id", "") or "EFARS"

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = KR_FONT
    style.font.size = Pt(10)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia"):
        rfonts.set(qn(attr), KR_FONT)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("전자금융사고 보고서 (별지 제2호서식 · 약식)")
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = ACCENT
    _apply_kr_font(run)

    mapping = _build_mapping(data)
    for label in [
        "보고일자", "보고종류", "case_id", "T0_KST", "사고유형", "사고등급",
        "등급확정권한", "보고기한_최초",
    ]:
        p = doc.add_paragraph()
        r = p.add_run(f"{label}: {mapping.get(label, PLACEHOLDER)}")
        r.font.size = Pt(10)
        _apply_kr_font(r)

    for sec_title, key in [
        ("1. 사고 개요 및 요약", "임원요약"),
        ("4. 결정 요청 사항", "결정요청"),
        ("5. 후속조치", "후속조치"),
    ]:
        h = doc.add_paragraph()
        hr = h.add_run(sec_title)
        hr.font.bold = True
        hr.font.size = Pt(12)
        hr.font.color.rgb = ACCENT
        _apply_kr_font(hr)
        _set_paragraph(doc.add_paragraph(), mapping.get(key, PLACEHOLDER))

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
